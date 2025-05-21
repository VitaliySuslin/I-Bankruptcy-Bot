import os
import logging
import base64
from datetime import datetime

from io import BytesIO

from PIL import Image
from telegram import Update
from telegram.ext import (
    ApplicationBuilder,
    CommandHandler,
    MessageHandler,
    ContextTypes,
    filters
)

from docx import Document
from docx import Document as DocxDocument
import PyPDF2

from src.bot.ai_client.client import ask_gpt
from src.bot.logger_setup import setup_logger
from src.bot.config import BOT_TOKEN


logger: logging.Logger = setup_logger()

def generate_word_doc(data: str) -> str:
    """
    Создаёт Word-документ с заявлением о банкротстве.

    :param data: Строка с текстом заявления
    :return: Путь к сохранённому файлу
    """
    logger.info("Начинаем генерацию Word-документа")

    doc = Document()
    doc.add_heading('Заявление о признании гражданина банкротом', level=1)

    paragraphs = data.split('\n')
    for para in paragraphs:
        stripped_para = para.strip()
        if stripped_para:
            doc.add_paragraph(stripped_para)

    # Генерация имени файла с текущим временем в формате ЧЧММСС
    timestamp = datetime.now().strftime("%H%M%S")
    file_name = f"Заявление_{timestamp}.docx"
    
    try:
        doc.save(file_name)
        logger.info(f"Word-документ успешно сохранён: {file_name}")
    except Exception as e:
        logger.error("Ошибка при сохранении Word-документа", exc_info=True)
        raise

    return file_name

async def start(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    logger.info("Команда /start получена")
    try:
        await update.message.reply_text("Привет! Я помогу вам подготовить заявление о банкротстве. Введите /begin чтобы начать.")
        logger.info("Сообщение по команде /start отправлено")
    except Exception as e:
        logger.error("Ошибка при обработке команды /start", exc_info=True)
        raise

async def begin(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    logger.info("Команда /begin получена")
    try:
        await update.message.reply_text("Пожалуйста, загрузите ваши документы: паспорт, справки о доходах, информацию о долгах и т.п. Это могут быть фото, PDF или Word-файлы.")
        logger.info("Сообщение по команде /begin отправлено")
    except Exception as e:
        logger.error("Ошибка при обработке команды /begin", exc_info=True)
        raise


async def handle_document(update: Update, context: ContextTypes.DEFAULT_TYPE):
    logger.info("Получен документ/фото от пользователя")
    msg = update.message

    try:
        if msg.document:
            logger.info("Обработка документа")
            file = await msg.document.get_file()
            file_name = msg.document.file_name
            file_type = file_name.split('.')[-1].lower()

            file_data = await file.download_as_bytearray()
            file_io = BytesIO(file_data)

            extracted_text = ""

            if file_type == 'pdf':
                reader = PyPDF2.PdfReader(file_io)
                extracted_text = "\n".join([page.extract_text() for page in reader.pages])

            elif file_type == 'docx':
                doc = DocxDocument(file_io)
                extracted_text = "\n".join([para.text for para in doc.paragraphs])

            elif file_type == 'txt':
                extracted_text = file_io.getvalue().decode('utf-8')

            else:
                await msg.reply_text("Поддерживаются только форматы: PDF, DOCX, TXT.")
                logger.warning("Пользователь загрузил неподдерживаемый формат файла")
                return

            await msg.reply_text("Обрабатываю документ...")

            prompt = f"""
                На основании следующего текста из документа:
                {extracted_text[:3000]}
                Извлеки анкетные данные заявителя и информацию о долгах:
                - ФИО (если есть)
                - Дата рождения (если есть)
                - Адрес регистрации (если есть)
                - Паспортные данные (если есть)
                - Общая сумма долга (если есть)
                - Кредиторы (если есть)
            """

            data_for_gpt = {
                "type": "text",
                "content": prompt
            }

            extracted_data = await ask_gpt(data_for_gpt)
            await generate_and_send_doc(extracted_data.choices[0].message.content, update)

        elif msg.photo:
            logger.info("Обработка фотографии")
            photo = msg.photo[-1]
            file = await photo.get_file()
            image_data = await file.download_as_bytearray()
            image = Image.open(BytesIO(image_data))
            buffered = BytesIO()
            image.save(buffered, format=image.format)
            img_str = base64.b64encode(buffered.getvalue()).decode("utf-8")

            await msg.reply_text("Обрабатываю изображение...")

            data_for_gpt = {
                "type": "image",
                "content": img_str,
                "mime_type": f"image/{image.format.lower()}"
            }

            extracted_data = await ask_gpt(data_for_gpt)
            await generate_and_send_doc(extracted_data.choices[0].message.content, update)

    except Exception as e:
        logger.exception("Ошибка при обработке документа/фото")
        await msg.reply_text("Произошла ошибка при обработке вашего документа.")


async def generate_and_send_doc(data: str, update: Update):
    logger.info("Начинаем генерацию заявления")
    try:
        claim_prompt = f"""
            На основании следующих данных:
            {data}
            
            Составь официальное заявление в суд о признании гражданина банкротом согласно ФЗ №127-ФЗ.
        """

        data_for_gpt = {
            "type": "text",
            "content": claim_prompt
        }

        claim_text_response = await ask_gpt(data_for_gpt)
        claim_text = claim_text_response.choices[0].message.content

        file_path = generate_word_doc(claim_text)

        await update.message.reply_document(document=open(file_path, 'rb'))
        os.remove(file_path)
        logger.info(f"Документ отправлен пользователю и удалён: {file_path}")

    except Exception as e:
        logger.exception("Ошибка при генерации или отправке заявления")
        await update.message.reply_text("Не удалось сформировать заявление. Попробуйте позже.")


def main():
    logger.info("Запуск Telegram-бота")
    try:
        app = ApplicationBuilder().token(BOT_TOKEN).build()

        app.add_handler(CommandHandler("start", start))
        app.add_handler(CommandHandler("begin", begin))
        app.add_handler(MessageHandler(filters.Document.ALL | filters.PHOTO, handle_document))

        logger.info("Бот запущен и ожидает сообщений")
        app.run_polling()
    except Exception as e:
        logger.critical("Критическая ошибка при запуске бота", exc_info=True)
        raise